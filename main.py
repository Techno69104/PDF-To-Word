import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import os

def extract_text_with_pypdf2(pdf_path):
    """Extract text from PDF using PyPDF2"""
    text_by_page = []
    try:
        with open(pdf_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                text_by_page.append(text if text else "No text found on this page.")
        return text_by_page
    except Exception as e:
        print(f"Error extracting text with PyPDF2: {e}")
        return ["Error extracting text"] * 10  # Return placeholder

def extract_images_with_pymupdf(pdf_path):
    """Extract images from PDF using PyMuPDF"""
    images_by_page = {}
    try:
        pdf_document = fitz.open(pdf_path)
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images()
            images = []
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image = Image.open(io.BytesIO(image_bytes))
                images.append(image)
            images_by_page[page_num + 1] = images
        pdf_document.close()
        return images_by_page
    except Exception as e:
        print(f"Error extracting images with PyMuPDF: {e}")
        return {}

def pdf_to_word(pdf_path, word_path):
    """Convert PDF to Word document with text and images"""
    try:
        # Extract text from PDF
        print(f"Extracting text from: {pdf_path}")
        text_by_page = extract_text_with_pypdf2(pdf_path)
        
        # Extract images from PDF
        print("Extracting images from PDF...")
        images_by_page = extract_images_with_pymupdf(pdf_path)
        
        # Create a Word document
        print("Creating Word document...")
        doc = Document()
        
        # Add content page by page
        for page_num, text in enumerate(text_by_page, start=1):
            # Add page header
            doc.add_heading(f'Page {page_num}', level=1)
            
            # Add text content
            doc.add_paragraph(text)
            
            # Add images from this page
            if page_num in images_by_page and images_by_page[page_num]:
                doc.add_heading('Images on this page:', level=2)
                for img_index, img in enumerate(images_by_page[page_num]):
                    # Save image temporarily to add to Word
                    temp_img_path = f"temp_img_{page_num}_{img_index}.png"
                    img.save(temp_img_path)
                    doc.add_picture(temp_img_path, width=Inches(4))
                    # Clean up temp file
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
            
            # Add page break except after last page
            if page_num < len(text_by_page):
                doc.add_page_break()
        
        # Save the document
        doc.save(word_path)
        print(f"Word document saved to: {word_path}")
        return True
        
    except Exception as e:
        print(f"Error in pdf_to_word: {e}")
        raise e

# For command line usage (backward compatibility)
if __name__ == "__main__":
    # This runs only when script is executed directly
    input_pdf = "input_pdf_file.pdf"
    output_word = "output_word_file.docx"
    
    if os.path.exists(input_pdf):
        pdf_to_word(input_pdf, output_word)
        print(f"Conversion complete! Output saved to {output_word}")
    else:
        print(f"Error: {input_pdf} not found!")
        print("Please place your PDF file in the project directory and name it 'input_pdf_file.pdf'")
