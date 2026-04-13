"""
PDF to Word Converter with OCR Support & Image Preprocessing
Optimized for CamScanner and other scanned PDFs
"""

import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageEnhance, ImageFilter
import io
import os
import re
import sys
import tempfile

# Try to import OCR libraries
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    print("OCR libraries loaded successfully", file=sys.stderr)
except ImportError as e:
    OCR_AVAILABLE = False
    print(f"Warning: OCR libraries not installed. Error: {e}", file=sys.stderr)

def preprocess_image_for_ocr(image):
    """
    Optimize CamScanner images for better OCR accuracy
    CamScanner often produces: low contrast, slightly rotated, grayscale images
    """
    try:
        # Convert to RGB if needed
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Convert to grayscale for OCR
        gray_image = image.convert('L')
        
        # Increase contrast (CamScanner often has low contrast)
        contrast_enhancer = ImageEnhance.Contrast(gray_image)
        gray_image = contrast_enhancer.enhance(2.0)
        
        # Apply sharpening filter
        gray_image = gray_image.filter(ImageFilter.SHARPEN)
        
        # Apply unsharp mask for edge enhancement
        gray_image = gray_image.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
        
        # Binarize image (convert to pure black and white) for better OCR
        threshold = 150
        binary_image = gray_image.point(lambda p: 255 if p > threshold else 0, '1')
        
        # Convert back to RGB for pytesseract
        final_image = binary_image.convert('RGB')
        
        # Resize if too small (minimum 300 DPI equivalent)
        width, height = final_image.size
        if width < 1000:
            ratio = 1000 / width
            new_size = (int(width * ratio), int(height * ratio))
            final_image = final_image.resize(new_size, Image.Resampling.LANCZOS)
        
        return final_image
        
    except Exception as e:
        print(f"Image preprocessing error: {e}", file=sys.stderr)
        return image  # Return original if preprocessing fails

def extract_text_with_pypdf2(pdf_path):
    """Extract text from digital PDFs using PyPDF2"""
    text_by_page = []
    try:
        with open(pdf_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text = re.sub(r'\s+', ' ', text)
                    text_by_page.append(text.strip())
                else:
                    text_by_page.append("")
        return text_by_page
    except Exception as e:
        print(f"PyPDF2 error: {e}", file=sys.stderr)
        return []

def extract_text_with_ocr_enhanced(pdf_path):
    """
    Extract text from scanned PDFs using OCR with preprocessing
    Optimized for CamScanner documents
    """
    if not OCR_AVAILABLE:
        print("OCR not available - libraries missing", file=sys.stderr)
        return None
    
    text_by_page = []
    try:
        print(f"Starting enhanced OCR on: {pdf_path}", file=sys.stderr)
        
        # Convert PDF pages to high-resolution images
        # Use 400 DPI for better quality (CamScanner works better at higher DPI)
        images = convert_from_path(pdf_path, dpi=400)
        print(f"Converted {len(images)} pages to images at 400 DPI", file=sys.stderr)
        
        for page_num, image in enumerate(images):
            print(f"Preprocessing page {page_num + 1} for OCR...", file=sys.stderr)
            
            # Apply CamScanner-specific preprocessing
            processed_image = preprocess_image_for_ocr(image)
            
            # Save preprocessed image temporarily for debugging (optional)
            # processed_image.save(f"/tmp/preprocessed_page_{page_num + 1}.png")
            
            print(f"Running OCR on page {page_num + 1}...", file=sys.stderr)
            
            # Configure Tesseract for better accuracy with scanned documents
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?;:()[]{}<>/\\|@#$%^&*+=_- '
            
            # Try with default config first
            text = pytesseract.image_to_string(processed_image, config=custom_config)
            
            # If result is poor, try with different page segmentation mode
            if len(text.strip()) < 50:
                print(f"Page {page_num + 1}: Low text yield, trying alternative PSM mode...", file=sys.stderr)
                alt_config = r'--oem 3 --psm 3'  # Fully automatic page segmentation
                text = pytesseract.image_to_string(processed_image, config=alt_config)
            
            if text and text.strip():
                # Clean up OCR text
                text = re.sub(r'\s+', ' ', text)
                text = text.strip()
                text_by_page.append(text)
                print(f"Page {page_num + 1}: Extracted {len(text)} characters", file=sys.stderr)
            else:
                print(f"Page {page_num + 1}: No text found after OCR", file=sys.stderr)
                text_by_page.append("[No readable text found on this page. The document may be handwritten or have poor image quality.]")
        
        return text_by_page if text_by_page else None
        
    except Exception as e:
        print(f"OCR error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return None

def extract_images_with_pymupdf(pdf_path):
    """Extract images from PDF using PyMuPDF"""
    images_by_page = {}
    try:
        pdf_document = fitz.open(pdf_path)
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images()
            images = []
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Preprocess extracted images as well (if they are scanned content)
                    if image.width > 500 and image.height > 500:
                        image = preprocess_image_for_ocr(image)
                    
                    images.append(image)
                except Exception as img_error:
                    print(f"Error extracting image {img_index}: {img_error}", file=sys.stderr)
            images_by_page[page_num + 1] = images
        pdf_document.close()
        return images_by_page
    except Exception as e:
        print(f"Image extraction error: {e}", file=sys.stderr)
        return {}

def clean_ocr_text(text):
    """Clean and improve OCR extracted text"""
    if not text:
        return ""
    
    # Remove excessive spaces
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common OCR issues in CamScanner
    replacements = {
        r'\|': 'I',
        r'0(?=\d)': 'O',
        r'1(?=\s)': 'I',
        r'’': "'",
        r'‘': "'",
        r'“': '"',
        r'”': '"',
        r'…': '...',
    }
    
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text)
    
    # Fix common word spacing issues
    text = re.sub(r'(\w)\s+(\W)', r'\1\2', text)
    text = re.sub(r'(\W)\s+(\w)', r'\1 \2', text)
    
    return text.strip()

def pdf_to_word(pdf_path, word_path):
    """
    Convert PDF to Word document with enhanced preprocessing
    Optimized for CamScanner and scanned PDFs
    """
    try:
        print(f"Processing PDF: {pdf_path}", file=sys.stderr)
        
        # Check if file exists
        if not os.path.exists(pdf_path):
            raise Exception(f"PDF file not found: {pdf_path}")
        
        # First try regular text extraction (for digital PDFs)
        text_by_page = extract_text_with_pypdf2(pdf_path)
        
        # Check if we got meaningful text
        has_text = False
        if text_by_page and len(text_by_page) > 0:
            first_page_text = text_by_page[0] if text_by_page[0] else ""
            has_text = len(first_page_text) > 100
            print(f"Digital text extraction: found {len(first_page_text)} chars on page 1", file=sys.stderr)
        
        # If no text found or very little text, use enhanced OCR
        if not has_text:
            print("No selectable text found or insufficient text. Using enhanced OCR with preprocessing...", file=sys.stderr)
            
            if OCR_AVAILABLE:
                # Use enhanced OCR with preprocessing
                ocr_text = extract_text_with_ocr_enhanced(pdf_path)
                
                if ocr_text and len(ocr_text) > 0:
                    text_by_page = ocr_text
                    # Clean OCR text
                    text_by_page = [clean_ocr_text(text) for text in text_by_page]
                    
                    # Calculate total extracted characters
                    total_chars = sum(len(text) for text in text_by_page if text)
                    print(f"Enhanced OCR successful! Extracted {len(text_by_page)} pages, {total_chars} total characters", file=sys.stderr)
                    
                    # If still no text, try alternative approach
                    if total_chars < 50:
                        print("Low text yield, attempting secondary OCR method...", file=sys.stderr)
                        text_by_page = attempt_alternative_ocr(pdf_path, text_by_page)
                else:
                    print("Enhanced OCR returned no text", file=sys.stderr)
                    text_by_page = ["[This appears to be a scanned PDF. OCR processing could not extract readable text. For best results with CamScanner PDFs, try: 1) Upload to Google Drive → Open with Google Docs, or 2) Ensure the original scan has good contrast and clarity.]"]
            else:
                print("OCR not available", file=sys.stderr)
                text_by_page = ["[This appears to be a scanned PDF. OCR support requires additional server configuration. Please contact support or use Google Drive to convert this document.]"]
        
        # Ensure text_by_page is a list with at least one element
        if not text_by_page or not isinstance(text_by_page, list) or len(text_by_page) == 0:
            text_by_page = ["[No text could be extracted from this PDF. The file may be empty or corrupted.]"]
        
        # Extract images (optional)
        images_by_page = extract_images_with_pymupdf(pdf_path)
        
        # Create Word document
        print("Creating Word document...", file=sys.stderr)
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Add cover page with metadata
        doc.add_heading('PDF to Word Conversion Result', 0)
        doc.add_paragraph(f'Converted on: {__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total pages: {len(text_by_page)}')
        doc.add_paragraph('Source: Scanned PDF (CamScanner optimized)')
        doc.add_page_break()
        
        # Add content page by page
        for page_num in range(len(text_by_page)):
            # Add page header
            heading = doc.add_heading(f'Page {page_num + 1}', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add text content
            if page_num < len(text_by_page) and text_by_page[page_num]:
                text = text_by_page[page_num]
                
                # Skip placeholder text that indicates no content
                if len(text) > 10 and not text.startswith("[No"):
                    # Split long text into paragraphs by sentences or line breaks
                    # First try splitting by double newlines
                    if '\n\n' in text:
                        paragraphs = text.split('\n\n')
                    else:
                        # Split by periods for better readability
                        import re
                        paragraphs = re.split(r'(?<=[.!?])\s+', text)
                    
                    for para in paragraphs:
                        if para.strip():
                            paragraph = doc.add_paragraph()
                            # Remove excessive spaces
                            clean_para = re.sub(r'\s+', ' ', para.strip())
                            run = paragraph.add_run(clean_para)
                            run.font.size = Pt(11)
                else:
                    # Add placeholder text
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph("[No text extracted from this page]")
            
            # Add images from this page (limit to 3 images per page)
            if page_num + 1 in images_by_page and images_by_page[page_num + 1]:
                images = images_by_page[page_num + 1][:3]
                if images:
                    doc.add_paragraph()  # Add spacing
                    doc.add_heading('Images on this page:', level=2)
                    
                    for img_index, img in enumerate(images):
                        try:
                            # Resize large images for Word document
                            if img.width > 500:
                                ratio = 500 / img.width
                                new_width = 500
                                new_height = int(img.height * ratio)
                                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                            
                            # Save image temporarily
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                                img.save(temp_img.name, 'PNG')
                                doc.add_picture(temp_img.name, width=Inches(4))
                                os.unlink(temp_img.name)
                        except Exception as img_error:
                            print(f"Error adding image: {img_error}", file=sys.stderr)
            
            # Add page break except after last page
            if page_num < len(text_by_page) - 1:
                doc.add_page_break()
        
        # Save the document
        doc.save(word_path)
        
        # Verify file was created
        if os.path.exists(word_path) and os.path.getsize(word_path) > 0:
            print(f"Word document saved successfully: {word_path} (Size: {os.path.getsize(word_path)} bytes)", file=sys.stderr)
            return True
        else:
            raise Exception("Output file is empty or was not created")
        
    except Exception as e:
        print(f"Error in pdf_to_word: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        raise e

def attempt_alternative_ocr(pdf_path, current_text_by_page):
    """Attempt alternative OCR method if primary fails"""
    try:
        print("Attempting alternative OCR method with different settings...", file=sys.stderr)
        
        # Try with lower DPI and different preprocessing
        images = convert_from_path(pdf_path, dpi=300)
        alternative_text = []
        
        for page_num, image in enumerate(images):
            # Minimal preprocessing
            if image.mode != 'L':
                image = image.convert('L')
            
            # Try with different Tesseract configuration
            config = r'--oem 1 --psm 4'  # Assume a single column of text
            text = pytesseract.image_to_string(image, config=config)
            
            if text and text.strip():
                text = re.sub(r'\s+', ' ', text).strip()
                alternative_text.append(text)
            else:
                alternative_text.append(current_text_by_page[page_num] if page_num < len(current_text_by_page) else "[No text found]")
        
        return alternative_text
    except Exception as e:
        print(f"Alternative OCR failed: {e}", file=sys.stderr)
        return current_text_by_page

# For command line testing
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        input_pdf = sys.argv[1]
        output_word = sys.argv[2] if len(sys.argv) > 2 else "output.docx"
        pdf_to_word(input_pdf, output_word)
        print(f"Conversion complete! Output saved to {output_word}")
